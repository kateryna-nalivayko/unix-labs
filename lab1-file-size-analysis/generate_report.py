"""Generate lab report as .docx file with actual analysis results and plot images."""

import os
import sys

# Ensure we can import core modules
from core.settings.conf_reader import SETTINGS
from core.analyzer import load_data, compute_stats, find_dominant_range

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def fmt_size(n):
    if n < 1024:
        return f"{n} B"
    elif n < 1024 ** 2:
        return f"{n / 1024:.1f} KB"
    elif n < 1024 ** 3:
        return f"{n / 1024 ** 2:.1f} MB"
    else:
        return f"{n / 1024 ** 3:.2f} GB"


def generate_report(data_path="./data/result.txt", plots_dir="./plots", output_path="./report.docx"):
    if not os.path.exists(data_path):
        print(f"Error: {data_path} not found. Run 'make collect' first.")
        sys.exit(1)

    # Load data and compute stats
    sizes = load_data(data_path)
    stats = compute_stats(sizes)
    dominant = find_dominant_range(sizes)

    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # -- Title --
    title = doc.add_heading("", level=0)
    run = title.add_run("Лабораторна робота №1\nЧастотний аналіз розподілу розмірів файлів")
    run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -- 1. Опис завдання --
    doc.add_heading("1. Опис завдання", level=1)
    doc.add_paragraph(
        "Мета роботи — проаналізувати частотний розподіл розмірів файлів "
        "у файловій системі UNIX (macOS). Необхідно зібрати дані про розміри "
        "всіх файлів у файловій системі, побудувати візуалізації розподілу "
        "(гістограма, CDF, boxplot) та сформулювати висновки про характер розподілу."
    )

    # -- 2. Підготовка даних --
    doc.add_heading("2. Підготовка даних", level=1)
    doc.add_paragraph(
        "Збір даних здійснюється рекурсивним обходом файлової системи починаючи "
        "з кореневого каталогу / за допомогою функції os.scandir() мови Python. "
        "Для кожного файлу зчитується його розмір через системний виклик stat(). "
        "Результати записуються у файл data/result.txt — один розмір (у байтах) на рядок."
    )
    doc.add_paragraph("Обробка помилок:")
    for item in [
        "Символічні посилання пропускаються",
        "Помилки доступу (PermissionError) ігноруються",
        "Інші помилки ОС (OSError) ігноруються",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Еквівалент одним рядком у shell:")
    p = doc.add_paragraph()
    run = p.add_run("find / -type f -exec stat -f '%z' {} + 2>/dev/null > data/result.txt")
    run.font.name = "Courier New"
    run.font.size = Pt(11)

    doc.add_paragraph(f"Загалом зібрано {stats['count']:,} файлів.")

    # -- 3. Архітектура рішення --
    doc.add_heading("3. Архітектура рішення", level=1)
    doc.add_paragraph(
        "Рішення побудовано за принципом розділення відповідальності (separation of concerns):"
    )
    for item in [
        "collector — рекурсивний обхід файлової системи, стрімінг розмірів у файл",
        "analyzer — завантаження даних, обчислення статистичних метрик, формулювання висновків",
        "visualizer — побудова графіків (гістограма, CDF, boxplot)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Конвеєр обробки: collector (os.scandir) → data/result.txt → analyzer (numpy) → plots/ + висновки")

    # -- 4. Аналіз та візуалізація --
    doc.add_heading("4. Аналіз та візуалізація", level=1)

    doc.add_heading("4.1 Статистичні метрики", level=2)

    table = doc.add_table(rows=1, cols=2, style="Table Grid")
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Метрика"
    hdr_cells[1].text = "Значення"
    for hdr in hdr_cells:
        for p in hdr.paragraphs:
            p.runs[0].bold = True

    metrics = [
        ("Кількість файлів", f"{stats['count']:,}"),
        ("Мінімальний розмір", fmt_size(stats["min"])),
        ("Максимальний розмір", fmt_size(stats["max"])),
        ("Середнє арифметичне", fmt_size(int(stats["mean"]))),
        ("Медіана", fmt_size(int(stats["median"]))),
        ("Стандартне відхилення", fmt_size(int(stats["std"]))),
    ]
    for pctile in [50, 75, 90, 95, 99]:
        metrics.append((f"Перцентиль P{pctile}", fmt_size(int(stats["percentiles"][pctile]))))

    for label, value in metrics:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    doc.add_paragraph()

    doc.add_heading("4.2 Розподіл за діапазонами розмірів", level=2)

    table2 = doc.add_table(rows=1, cols=3, style="Table Grid")
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Діапазон"
    hdr2[1].text = "Кількість"
    hdr2[2].text = "Відсоток"
    for h in hdr2:
        for p in h.paragraphs:
            p.runs[0].bold = True

    for label, count, pct in stats["buckets"]:
        row = table2.add_row().cells
        row[0].text = label
        row[1].text = f"{count:,}"
        row[2].text = f"{pct:.1f}%"

    doc.add_paragraph()

    doc.add_heading("4.3 Візуалізація", level=2)

    # Histogram
    hist_path = os.path.join(plots_dir, "histogram.png")
    if os.path.exists(hist_path):
        doc.add_paragraph("Гістограма розподілу розмірів файлів (логарифмічна шкала по осі X):")
        doc.add_picture(hist_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # CDF
    cdf_path = os.path.join(plots_dir, "cdf.png")
    if os.path.exists(cdf_path):
        doc.add_paragraph()
        doc.add_paragraph("Кумулятивна функція розподілу (CDF) з позначеними перцентилями:")
        doc.add_picture(cdf_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Boxplot
    box_path = os.path.join(plots_dir, "boxplot.png")
    if os.path.exists(box_path):
        doc.add_paragraph()
        doc.add_paragraph("Діаграма розмахів (boxplot) на логарифмічній шкалі:")
        doc.add_picture(box_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -- 5. Висновки --
    doc.add_heading("5. Висновки", level=1)

    doc.add_paragraph(
        f"За результатами аналізу {stats['count']:,} файлів файлової системи macOS "
        f"було встановлено наступне:"
    )

    for pct, lo, hi in dominant:
        doc.add_paragraph(
            f"{pct}% файлів мають розміри у діапазоні від {fmt_size(int(lo))} до {fmt_size(int(hi))}",
            style="List Bullet",
        )

    best = dominant[-1]
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        f"Переважна більшість файлів ({best[0]}%) має розміри у діапазоні "
        f"від {fmt_size(int(best[1]))} до {fmt_size(int(best[2]))}."
    )
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        "Розподіл розмірів файлів має характерну правосторонню асиметрію "
        "(right-skewed distribution): медіана значно менша за середнє арифметичне, "
        f"що свідчить про наявність великої кількості малих файлів та відносно "
        f"невеликої кількості дуже великих файлів. "
        f"Медіана становить {fmt_size(int(stats['median']))}, тоді як середнє — "
        f"{fmt_size(int(stats['mean']))}, що відрізняється на порядки."
    )

    # -- 6. Як запустити --
    doc.add_heading("6. Як запустити", level=1)
    commands = [
        ("Встановлення залежностей:", "uv sync"),
        ("Збір даних:", "make collect"),
        ("Аналіз та візуалізація:", "make analyze"),
        ("Генерація цього звіту:", "uv run --with python-docx python generate_report.py"),
    ]
    for desc, cmd in commands:
        doc.add_paragraph(desc)
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.font.name = "Courier New"
        run.font.size = Pt(11)

    doc.save(output_path)
    print(f"Report saved to {output_path}")


if __name__ == "__main__":
    generate_report()
